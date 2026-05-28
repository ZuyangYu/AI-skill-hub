#!/usr/bin/env python
"""Scrape ecommerce product pages into JSON, images, and optional DOCX output."""

from __future__ import annotations

import argparse
import asyncio
import json
import os
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse


IMAGE_MIN_BYTES = 5_000


def is_supported_image(body: bytes) -> bool:
    return (
        body.startswith(b"\xff\xd8\xff")
        or body.startswith(b"\x89PNG\r\n\x1a\n")
        or (len(body) > 12 and body[:4] == b"RIFF" and body[8:12] == b"WEBP")
    )


def image_extension(body: bytes) -> str:
    if body.startswith(b"\xff\xd8\xff"):
        return ".jpg"
    if body.startswith(b"\x89PNG\r\n\x1a\n"):
        return ".png"
    if len(body) > 12 and body[:4] == b"RIFF" and body[8:12] == b"WEBP":
        return ".webp"
    return ".img"


@dataclass
class ScrapeConfig:
    url: str
    site: str = "generic"
    output_root: Path = Path("scraped_data")
    outputs: set[str] = field(default_factory=lambda: {"json", "images"})
    headed: bool = False
    timeout_ms: int = 90_000
    image_width: int = 1400
    max_colors: int | None = None
    update_index: bool = True


def slugify(value: str, fallback: str = "product") -> str:
    value = re.sub(r"[^A-Za-z0-9._-]+", "-", value.strip()).strip("-._")
    return value[:80] or fallback


def normalize_space(value: str | None) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def extract_product_id(url: str, text: str = "", site: str = "generic") -> str:
    parsed = urlparse(url)
    query = parse_qs(parsed.query)

    if site == "patagonia":
        for values in query.values():
            for value in values:
                match = re.search(r"\b(\d{4,8})\b", value)
                if match:
                    return match.group(1)
        match = re.search(r"Style No\.\s*(\d{4,8})", text)
        if match:
            return match.group(1)
        match = re.search(r"/(\d{4,8})(?:[/?#._-]|$)", parsed.path)
        if match:
            return match.group(1)

    path_bits = [bit for bit in parsed.path.split("/") if bit]
    if path_bits:
        return slugify(path_bits[-1])
    return slugify(parsed.netloc)


def merge_missing(target: dict[str, Any], source: dict[str, Any]) -> None:
    for key, value in source.items():
        if value not in (None, "", [], {}) and target.get(key) in (None, "", [], {}):
            target[key] = value


async def launch_page(config: ScrapeConfig):
    from playwright.async_api import async_playwright

    playwright = await async_playwright().start()
    browser = await playwright.chromium.launch(
        headless=not config.headed,
        args=["--disable-blink-features=AutomationControlled"],
    )
    context = await browser.new_context(
        user_agent=(
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/125.0.0.0 Safari/537.36"
        ),
        viewport={"width": 1920, "height": 1080},
    )
    page = await context.new_page()
    await page.add_init_script(
        'Object.defineProperty(navigator, "webdriver", {get: () => undefined});'
    )
    return playwright, browser, context, page


async def extract_page_data(page, config: ScrapeConfig) -> dict[str, Any]:
    if config.site == "patagonia":
        await expand_patagonia_details(page)

    data = await page.evaluate(
        """() => {
        const text = document.body ? document.body.innerText : "";
        const meta = (name) => {
          const el = document.querySelector(
            `meta[property="${name}"], meta[name="${name}"]`
          );
          return el ? el.getAttribute("content") || "" : "";
        };

        const jsonLd = [];
        document.querySelectorAll('script[type="application/ld+json"]').forEach((el) => {
          try {
            const parsed = JSON.parse(el.textContent || "{}");
            if (Array.isArray(parsed)) jsonLd.push(...parsed);
            else jsonLd.push(parsed);
          } catch (e) {}
        });

        const images = [];
        document.querySelectorAll("img, source").forEach((el) => {
          const candidates = [
            el.currentSrc,
            el.src,
            el.getAttribute("src"),
            el.getAttribute("data-src"),
            el.getAttribute("data-zoom"),
            el.getAttribute("data-hires")
          ].filter(Boolean);
          const srcset = el.getAttribute("srcset") || el.getAttribute("data-srcset") || "";
          srcset.split(",").forEach((part) => {
            const url = part.trim().split(/\\s+/)[0];
            if (url) candidates.push(url);
          });
          candidates.forEach((url) => images.push({
            url,
            alt: el.getAttribute("alt") || el.getAttribute("aria-label") || ""
          }));
        });

        const colorNames = {};
        document.querySelectorAll('[aria-label*="color" i], [data-color]').forEach((el) => {
          const code = el.getAttribute("data-color") || el.closest("[data-color]")?.getAttribute("data-color");
          const label = el.getAttribute("aria-label") || "";
          if (code && code !== "000") colorNames[code] = label.replace(/^color\\s+/i, "").trim();
        });

        return {
          title: document.querySelector("h1")?.innerText || document.title || "",
          text,
          meta: {
            title: meta("og:title") || meta("twitter:title"),
            description: meta("description") || meta("og:description"),
            image: meta("og:image") || meta("twitter:image"),
          price: meta("product:price:amount") || meta("og:price:amount"),
            currency: meta("product:price:currency")
          },
          jsonLd,
          images,
          colorNames
        };
    }"""
    )

    product: dict[str, Any] = {
        "source_url": config.url,
        "site": config.site,
        "product_id": "",
        "name": normalize_space(data.get("title") or data["meta"].get("title")),
        "brand": "",
        "price": "",
        "currency": data["meta"].get("currency", ""),
        "style_number": "",
        "color_code": "",
        "color_names": data.get("colorNames", {}),
        "description": normalize_space(data["meta"].get("description")),
        "features": [],
        "materials": "",
        "care_instructions": [],
        "certifications": [],
        "variants": [],
        "images": [],
        "raw": {"meta": data.get("meta"), "json_ld": data.get("jsonLd")},
    }

    apply_json_ld(product, data.get("jsonLd", []))
    apply_text_heuristics(product, data.get("text", ""), config.site)

    product["product_id"] = extract_product_id(
        config.url, data.get("text", ""), config.site
    )
    if not product["style_number"]:
        product["style_number"] = product["product_id"] if product["product_id"].isdigit() else ""

    seen = set()
    for image in data.get("images", []):
        url = image.get("url", "")
        if not url or url.startswith("data:") or url in seen:
            continue
        seen.add(url)
        product["images"].append({"url": url, "alt": normalize_space(image.get("alt"))})
    if data["meta"].get("image") and data["meta"]["image"] not in seen:
        product["images"].insert(0, {"url": data["meta"]["image"], "alt": "primary"})

    return product


def apply_json_ld(product: dict[str, Any], json_ld: list[dict[str, Any]]) -> None:
    candidates: list[dict[str, Any]] = []
    for item in json_ld:
        if "@graph" in item and isinstance(item["@graph"], list):
            candidates.extend(x for x in item["@graph"] if isinstance(x, dict))
        candidates.append(item)

    for item in candidates:
        item_type = item.get("@type", "")
        if isinstance(item_type, list):
            is_product = "Product" in item_type or "ProductGroup" in item_type
        else:
            is_product = item_type in {"Product", "ProductGroup"}
        if not is_product:
            continue

        brand = item.get("brand", "")
        if isinstance(brand, dict):
            brand = brand.get("name", "")
        offers = item.get("offers", {})
        if isinstance(offers, list):
            offers = offers[0] if offers else {}

        merge_missing(
            product,
            {
                "name": item.get("name"),
                "brand": brand,
                "description": item.get("description"),
                "price": str(offers.get("price", "")) if isinstance(offers, dict) else "",
                "currency": offers.get("priceCurrency", "") if isinstance(offers, dict) else "",
            },
        )
        if item_type == "ProductGroup" and item.get("productGroupID"):
            product["product_id"] = str(item["productGroupID"])
            product["style_number"] = str(item["productGroupID"])
            colors = item.get("color") or []
            if isinstance(colors, list):
                product["raw"]["jsonld_colors"] = colors
        elif item.get("sku") or item.get("mpn"):
            sku = str(item.get("sku") or item.get("mpn"))
            color_code = sku.split("-")[-1] if "-" in sku else ""
            variant = {
                "sku": sku,
                "color_code": color_code,
                "color": item.get("color", ""),
                "price": str(offers.get("price", "")) if isinstance(offers, dict) else "",
                "currency": offers.get("priceCurrency", "") if isinstance(offers, dict) else "",
                "availability": offers.get("availability", "") if isinstance(offers, dict) else "",
                "url": offers.get("url", "") if isinstance(offers, dict) else "",
                "image": item.get("image") or (offers.get("image", "") if isinstance(offers, dict) else ""),
            }
            if not any(v.get("sku") == sku for v in product["variants"]):
                product["variants"].append(variant)
            if color_code and variant.get("color") and not product["color_names"].get(color_code):
                product["color_names"][color_code] = variant["color"]

    if product.get("color_code") and product.get("variants"):
        selected = next(
            (v for v in product["variants"] if v.get("color_code") == product["color_code"]),
            None,
        )
        if selected:
            merge_missing(
                product,
                {
                    "price": selected.get("price"),
                    "currency": selected.get("currency"),
                },
            )


def apply_text_heuristics(product: dict[str, Any], text: str, site: str) -> None:
    price_match = re.search(r"\$[\d,]+(?:\.\d{2})?", text)
    if price_match and not product.get("price"):
        product["price"] = price_match.group(0)

    if site == "patagonia":
        style_match = re.search(r"\b([A-Z0-9]{3,10})\s*\|\s*Style No\.\s*(\d{4,8})", text)
        if style_match:
            product["color_code"] = style_match.group(1)
            product["style_number"] = style_match.group(2)
            product["product_id"] = style_match.group(2)

        fit_match = re.search(r"Fit\s+(Regular [Ff]it|Slim [Ff]it|Relaxed [Ff]it)", text)
        weight_match = re.search(r"Weight\s+(\d+)\s*g\s*\(([^)]+)\)", text)
        origin_match = re.search(r"Country of Origin\s+(Made in [^\n.]+)", text)
        materials_match = re.search(
            r"Materials(?:\s*&\s*Care Instructions)?[\s\S]{0,1200}?(\d+(?:\.\d+)?-oz[^\n.]{0,220})",
            text,
            re.IGNORECASE,
        )
        care_match = re.search(
            r"(Machine Wash[^\n.]+(?:\.|$))",
            text,
            re.IGNORECASE,
        )
        features = extract_patagonia_features(text)

        facts = []
        if fit_match:
            facts.append({"name": "fit", "value": fit_match.group(1)})
        if weight_match:
            facts.append({"name": "weight", "value": f"{weight_match.group(1)} g ({weight_match.group(2)})"})
        elif weight_from_feature := re.search(r"\b(\d+)\s*g\s*\(([^)]+oz)\)", text):
            facts.append({"name": "weight", "value": f"{weight_from_feature.group(1)} g ({weight_from_feature.group(2)})"})
        if origin_match:
            facts.append({"name": "origin", "value": origin_match.group(1)})
        product["raw"]["facts"] = facts

        if materials_match and not product.get("materials"):
            product["materials"] = normalize_space(materials_match.group(1))
        if care_match:
            product["care_instructions"] = [
                normalize_space(part) for part in care_match.group(1).split(",") if normalize_space(part)
            ]
        if features and not product.get("features"):
            product["features"] = features
        if "Fair Trade Certified" in text:
            product["certifications"].append("Fair Trade Certified")
        if "bluesign" in text:
            product["certifications"].append("bluesign approved")


async def expand_patagonia_details(page) -> None:
    await page.evaluate(
        """() => {
        const labels = [
          "Specs & Features",
          "Materials",
          "Materials & Care Instructions",
          "Fit",
          "Reviews"
        ];
        for (const label of labels) {
          const candidates = [...document.querySelectorAll("button, summary, [role='button']")];
          for (const el of candidates) {
            const text = (el.innerText || el.getAttribute("aria-label") || "").trim();
            if (text.includes(label) && el.getAttribute("aria-expanded") !== "true") {
              try { el.click(); } catch (e) {}
            }
          }
        }
    }"""
    )
    await page.wait_for_timeout(1_000)


def extract_patagonia_features(text: str) -> list[dict[str, str]]:
    match = re.search(
        r"Specs\s*&\s*Features([\s\S]+?)(?:Materials\s*&\s*Care Instructions|Materials|Reviews|$)",
        text,
        re.IGNORECASE,
    )
    if not match:
        return []

    skip = {
        "Specs & Features",
        "Care Instructions",
        "Country of Origin",
        "Weight",
        "Fit",
        "Materials",
    }
    lines = [normalize_space(line) for line in match.group(1).splitlines()]
    lines = [line for line in lines if line and line not in skip]

    features: list[dict[str, str]] = []
    i = 0
    while i < len(lines):
        title = lines[i]
        if re.fullmatch(r"\d+\s*g\s*\([^)]+oz\)", title):
            i += 1
            continue
        if len(title) <= 90 and not re.search(r"[.!?]$", title):
            desc = ""
            if i + 1 < len(lines) and len(lines[i + 1]) > 40:
                desc = lines[i + 1]
                i += 1
            if title not in {f["title"] for f in features}:
                features.append({"title": title, "description": desc})
        i += 1

    return features[:12]


async def setup_patagonia_capture(page, config: ScrapeConfig, captured: dict[str, bytes]) -> None:
    expected_id = extract_product_id(config.url, site="patagonia")

    async def route_upgrade(route):
        url = route.request.url
        if "BDJB_PRD" in url and "sw=" in url:
            url = re.sub(r"sw=\d+", f"sw={config.image_width}", url)
            url = re.sub(r"sh=\d+", f"sh={config.image_width}", url)
            await route.continue_(url=url)
        else:
            await route.continue_()

    async def on_response(response):
        url = response.url
        if "BDJB_PRD" not in url or "hi-res" not in url:
            return
        if expected_id and expected_id not in url:
            return
        try:
            body = await response.body()
        except Exception:
            return
        if len(body) < IMAGE_MIN_BYTES or not is_supported_image(body):
            return
        match = re.search(r"/([^/?#]+)\.jpg", url)
        if not match:
            return
        name = slugify(match.group(1))
        if name not in captured or len(body) > len(captured[name]):
            captured[name] = body

    await page.route("**/dw/image/v2/BDJB_PRD/**", route_upgrade)
    page.on("response", lambda response: asyncio.create_task(on_response(response)))


async def discover_patagonia_colors(page) -> list[str]:
    colors = await page.evaluate(
        """() => {
        const codes = new Set();
        document.querySelectorAll("[data-color]").forEach((el) => {
          const code = el.getAttribute("data-color");
          if (code && code !== "000") codes.add(code);
        });
        const text = document.body ? document.body.innerText : "";
        for (const match of text.matchAll(/\\b([A-Z0-9]{3,10})\\s*\\|\\s*Style No\\./g)) {
          if (match[1] !== "000") codes.add(match[1]);
        }
        return [...codes];
    }"""
    )
    return [color for color in colors if color and color != "000"]


def patagonia_variant_url(url: str, style_number: str, color_code: str) -> str:
    base = re.sub(r"([?&])dwvar_[^=&]+_color=[^&#]+&?", r"\1", url).rstrip("?&")
    sep = "&" if "?" in base else "?"
    return f"{base}{sep}dwvar_{style_number}_color={color_code}"


async def scrape(config: ScrapeConfig) -> dict[str, Any]:
    playwright, browser, context, page = await launch_page(config)
    captured: dict[str, bytes] = {}

    try:
        if config.site == "patagonia":
            await setup_patagonia_capture(page, config, captured)

        await page.goto(config.url, wait_until="domcontentloaded", timeout=config.timeout_ms)
        await page.wait_for_timeout(4_000 if config.site == "generic" else 12_000)
        product = await extract_page_data(page, config)

        if config.site == "patagonia":
            colors = await discover_patagonia_colors(page)
            current = product.get("color_code")
            ordered = ([current] if current else []) + [c for c in colors if c != current]
            valid_colors = []
            variant_colors = ordered[1:]
            if config.max_colors is not None:
                variant_colors = variant_colors[: config.max_colors]
            for color in variant_colors:
                before = len(captured)
                await page.goto(
                    patagonia_variant_url(config.url, product.get("style_number") or product["product_id"], color),
                    wait_until="domcontentloaded",
                    timeout=60_000,
                )
                await page.wait_for_timeout(8_000)
                if len(captured) > before:
                    valid_colors.append(color)
            product["raw"]["discovered_colors"] = ordered
            product["raw"]["valid_colors"] = valid_colors

        product_dir = config.output_root / product["product_id"]
        image_dir = product_dir / "images"
        product_dir.mkdir(parents=True, exist_ok=True)
        if "images" in config.outputs:
            image_dir.mkdir(parents=True, exist_ok=True)
            save_captured_images(product, image_dir, captured)

        if "json" in config.outputs:
            (product_dir / "product.json").write_text(
                json.dumps(product, ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
        if "docx" in config.outputs:
            write_docx(product, product_dir / "product.docx", image_dir)
        if config.update_index:
            update_catalog_index(product, product_dir, image_dir, config.output_root)

        return product
    finally:
        await context.close()
        await browser.close()
        await playwright.stop()


def save_captured_images(product: dict[str, Any], image_dir: Path, captured: dict[str, bytes]) -> None:
    saved_images = []
    for name, body in captured.items():
        if not is_supported_image(body):
            continue
        filename = f"{name}{image_extension(body)}"
        path = image_dir / filename
        if not path.exists() or path.stat().st_size < len(body):
            path.write_bytes(body)
        saved_images.append({"filename": filename, "path": str(path), "source": "browser-response"})
    if saved_images:
        product["images"] = saved_images + product.get("images", [])


def update_catalog_index(product: dict[str, Any], product_dir: Path, image_dir: Path, output_root: Path) -> None:
    index_path = output_root / "catalog_index.json"
    now = datetime.now().astimezone().isoformat(timespec="seconds")

    index = {"updated_at": now, "products": []}
    if index_path.exists():
        try:
            index = json.loads(index_path.read_text(encoding="utf-8"))
            if not isinstance(index, dict):
                index = {"updated_at": now, "products": []}
            if not isinstance(index.get("products"), list):
                index["products"] = []
        except json.JSONDecodeError:
            index = {"updated_at": now, "products": []}

    local_images = [img for img in product.get("images", []) if img.get("filename")]
    variants = product.get("variants") or []
    variant_colors = [v.get("color_code") for v in variants if v.get("color_code")]
    raw_colors = product.get("raw", {}).get("discovered_colors") or product.get("raw", {}).get("jsonld_colors") or []
    colors = sorted(dict.fromkeys([c for c in [*variant_colors, *raw_colors] if c]))
    color_names = {k: v for k, v in (product.get("color_names") or {}).items() if v}
    category = extract_category(product.get("source_url", ""))

    record = {
        "product_id": product.get("product_id", ""),
        "style_number": product.get("style_number", ""),
        "name": product.get("name", ""),
        "brand": product.get("brand", ""),
        "site": product.get("site", ""),
        "source_url": product.get("source_url", ""),
        "category": category,
        "price": product.get("price", ""),
        "currency": product.get("currency", ""),
        "colors": colors,
        "color_names": color_names,
        "variant_count": len(variants),
        "image_count": len(local_images),
        "json_path": str(product_dir / "product.json"),
        "docx_path": str(product_dir / "product.docx"),
        "image_dir": str(image_dir),
        "keywords": build_keywords(product, category),
        "last_scraped_at": now,
    }

    products = [
        item
        for item in index.get("products", [])
        if item.get("product_id") != record["product_id"] or item.get("site") != record["site"]
    ]
    products.append(record)
    products.sort(key=lambda item: (item.get("site", ""), item.get("product_id", "")))
    index["updated_at"] = now
    index["products"] = products

    output_root.mkdir(parents=True, exist_ok=True)
    index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")


def extract_category(url: str) -> str:
    parsed = urlparse(url)
    query = parse_qs(parsed.query)
    if query.get("cgid"):
        return query["cgid"][0]
    bits = [bit for bit in parsed.path.split("/") if bit]
    if "shop" in bits:
        return "-".join(bits[bits.index("shop") + 1 :])
    return ""


def build_keywords(product: dict[str, Any], category: str) -> list[str]:
    values: list[str] = [
        product.get("product_id", ""),
        product.get("style_number", ""),
        product.get("name", ""),
        product.get("brand", ""),
        category,
    ]
    values.extend((product.get("color_names") or {}).keys())
    values.extend((product.get("color_names") or {}).values())
    for variant in product.get("variants") or []:
        values.extend([variant.get("sku", ""), variant.get("color_code", ""), variant.get("color", "")])

    keywords = []
    seen = set()
    for value in values:
        normalized = normalize_space(str(value)).lower()
        if normalized and normalized not in seen:
            seen.add(normalized)
            keywords.append(normalized)
    return keywords


def write_docx(product: dict[str, Any], output_path: Path, image_dir: Path) -> None:
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    doc = Document()
    doc.add_heading(product.get("name") or product.get("product_id") or "Product", level=1)
    for label, key in [
        ("Source", "source_url"),
        ("Brand", "brand"),
        ("Price", "price"),
        ("Style", "style_number"),
        ("Color", "color_code"),
        ("Materials", "materials"),
    ]:
        value = product.get(key)
        if value:
            doc.add_paragraph(f"{label}: {value}")

    if product.get("description"):
        doc.add_heading("Description", level=2)
        doc.add_paragraph(product["description"])

    for section, key in [
        ("Features", "features"),
        ("Care Instructions", "care_instructions"),
        ("Certifications", "certifications"),
    ]:
        values = product.get(key) or []
        if values:
            doc.add_heading(section, level=2)
            for value in values:
                doc.add_paragraph(str(value), style="List Bullet")

    local_images = [img for img in product.get("images", []) if img.get("filename")]
    if local_images:
        doc.add_heading("Images", level=2)
        for image in local_images:
            path = image_dir / image["filename"]
            if path.exists():
                doc.add_paragraph(image["filename"])
                try:
                    doc.add_picture(str(path), width=Inches(3.0))
                except Exception:
                    converted = path.with_suffix(path.suffix + ".png")
                    try:
                        with Image.open(path) as source:
                            source.convert("RGB").save(converted, "PNG")
                        doc.add_picture(str(converted), width=Inches(3.0))
                    finally:
                        if converted.exists():
                            converted.unlink()

    try:
        doc.save(output_path)
    except PermissionError:
        fallback = output_path.with_name(
            f"{output_path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{output_path.suffix}"
        )
        doc.save(fallback)


def parse_args(argv: list[str]) -> ScrapeConfig:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("url")
    parser.add_argument("--site", choices=["generic", "patagonia"], default="generic")
    parser.add_argument("--output-root", default="scraped_data")
    parser.add_argument("--outputs", nargs="+", choices=["json", "images", "docx"], default=["json", "images"])
    parser.add_argument("--headed", action="store_true")
    parser.add_argument("--timeout-ms", type=int, default=90_000)
    parser.add_argument("--image-width", type=int, default=1400)
    parser.add_argument(
        "--max-colors",
        type=int,
        default=None,
        help="Maximum additional color variants to visit; use 0 for a fast current-page smoke test.",
    )
    parser.add_argument(
        "--no-index",
        action="store_true",
        help="Do not update catalog_index.json after scraping.",
    )
    args = parser.parse_args(argv)
    return ScrapeConfig(
        url=args.url,
        site=args.site,
        output_root=Path(args.output_root),
        outputs=set(args.outputs),
        headed=args.headed,
        timeout_ms=args.timeout_ms,
        image_width=args.image_width,
        max_colors=args.max_colors,
        update_index=not args.no_index,
    )


def main(argv: list[str] | None = None) -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    config = parse_args(argv or sys.argv[1:])
    product = asyncio.run(scrape(config))
    print(json.dumps({
        "product_id": product.get("product_id"),
        "name": product.get("name"),
        "image_count": len(product.get("images", [])),
        "output_root": str(config.output_root / product.get("product_id", "")),
        "index_path": str(config.output_root / "catalog_index.json") if config.update_index else "",
    }, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
